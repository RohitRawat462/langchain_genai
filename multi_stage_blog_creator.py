import os
import streamlit as st
from dotenv import load_dotenv
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import Runnable
from langchain_core.output_parsers import StrOutputParser

from fpdf import FPDF
from docx import Document

# Load environment variables
load_dotenv()
groq_api_key = os.getenv("GROQ_API_KEY")

# Initialize LLM
llm = ChatGroq(api_key=groq_api_key, model_name="llama3-70b-8192")

# Prompts
outline_prompt = ChatPromptTemplate.from_messages([
    ("system", "You are a professional blog content strategist."),
    ("human", "Create a blog post outline for the topic: {topic}")
])

intro_prompt = ChatPromptTemplate.from_messages([
    ("system", "You are a professional blog writer."),
    ("human", "Write an introduction for this outline:\n{outline}")
])

section_prompt = ChatPromptTemplate.from_messages([
    ("system", "You are an expert content creator."),
    ("human", "Expand this outline into full blog sections:\n{outline}")
])

social_prompt = ChatPromptTemplate.from_messages([
    ("system", "You are a skilled social media marketer."),
    ("human", "Create a short Twitter thread or LinkedIn summary from this blog:\n{full_blog}")
])

# Parsers
parser = StrOutputParser()

# Chains
outline_chain: Runnable = outline_prompt | llm | parser
intro_chain: Runnable = intro_prompt | llm | parser
section_chain: Runnable = section_prompt | llm | parser
social_chain: Runnable = social_prompt | llm | parser

# Streamlit UI
st.set_page_config(page_title="📝 Blog Creator with Groq", layout="centered")
st.title("🚀 Multi-Stage Blog Creator")
st.write("Enter a topic and let AI write a blog with social media summary. Export in PDF, TXT, or DOCX!")

topic = st.text_input("Enter blog topic:", placeholder="e.g., Future of AI in Healthcare")
generate = st.button("✨ Generate Blog")

if generate and topic:
    with st.spinner("Generating blog..."):
        try:
            outline = outline_chain.invoke({"topic": topic})
            intro = intro_chain.invoke({"outline": outline})
            sections = section_chain.invoke({"outline": outline})
            full_blog = f"{intro}\n\n{sections}"
            social = social_chain.invoke({"full_blog": full_blog})

            st.subheader("📝 Blog Outline")
            st.markdown(outline)

            st.subheader("✍️ Introduction")
            st.markdown(intro)

            st.subheader("📚 Blog Content")
            st.markdown(sections)

            st.subheader("📣 Social Media Summary")
            st.markdown(social)

            # Create downloads
            blog_content = f"Title: {topic}\n\nOutline:\n{outline}\n\nIntro:\n{intro}\n\nContent:\n{sections}\n\nSocial Media Summary:\n{social}"

            # TXT download
            st.download_button("⬇️ Download as TXT", blog_content, file_name="blog.txt")

            # PDF download
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=12)
            
            for line in blog_content.split("\n"):
                safe_line = line.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 10, safe_line)
            pdf_output = "blog.pdf"
            pdf.output(pdf_output)
            with open(pdf_output, "rb") as f:
                st.download_button("⬇️ Download as PDF", f, file_name="blog.pdf")

            # DOCX download
            doc = Document()
            doc.add_heading(topic, 0)
            doc.add_heading("Outline", level=1)
            doc.add_paragraph(outline)
            doc.add_heading("Introduction", level=1)
            doc.add_paragraph(intro)
            doc.add_heading("Sections", level=1)
            doc.add_paragraph(sections)
            doc.add_heading("Social Media Summary", level=1)
            doc.add_paragraph(social)
            docx_output = "blog.docx"
            doc.save(docx_output)
            with open(docx_output, "rb") as f:
                st.download_button("⬇️ Download as DOCX", f, file_name="blog.docx")

        except Exception as e:
            st.error(f"❌ Error: {e}")
